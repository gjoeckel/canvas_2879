#!/usr/bin/env python3
"""
Convert Canvas HTML page to DOCX file using Pandoc.

This creates a DOCX file that structurally matches the HTML,
improving mapping accuracy for tracked changes.
"""

import argparse
import subprocess
import sys
from pathlib import Path
from bs4 import BeautifulSoup
import re

# Import path utilities
sys.path.insert(0, str(Path(__file__).parent.parent))
from utils.paths import (
    get_project_root,
    get_canvas_reference_docx,
    get_winter_updates_dir,
    get_winter_archive_dir
)

def extract_user_content(html_file_path):
    """Extract and clean the user_content div from HTML, including original-link if present."""
    with open(html_file_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'html.parser')
    user_content = soup.find('div', class_='user_content')

    if not user_content:
        raise ValueError("Could not find .user_content div in HTML file")

    # Check if original-link is already inside user_content (moved to top)
    original_link = user_content.find('div', class_='original-link')

    # If not found inside user_content, check outside (old location)
    if not original_link:
        # Look for it as a sibling of user_content
        content_div = soup.find('div', class_='content')
        if content_div:
            original_link = content_div.find('div', class_='original-link')
            # If found outside, move it to the top of user_content
            if original_link:
                original_link.extract()
                user_content.insert(0, original_link)

    # Clean the content (with image path handling - no downloads, only local files)
    cleaned_content = clean_html_content(user_content, html_file_path)

    return cleaned_content

def clean_html_content(element, html_file_path=None):
    """Clean HTML content for DOCX conversion, preserving formatting and images."""
    # Create a copy to avoid modifying the original
    cleaned = BeautifulSoup(str(element), 'html.parser')

    # Replace iframes with placeholder text
    for iframe in cleaned.find_all('iframe'):
        title = iframe.get('title', 'Video')
        placeholder = cleaned.new_tag('p')
        placeholder.string = f"[Video: {title}]"
        iframe.replace_with(placeholder)

    # Handle images - use existing local files with Canvas IDs only (no fallbacks)
    if html_file_path:
        html_file = Path(html_file_path)
        html_dir = html_file.parent

        for img in cleaned.find_all('img'):
            orig_src = img.get('src', '')
            if not orig_src:
                # No src, replace with "image not found"
                img.replace_with("image not found")
                continue

            # Get Canvas file ID from src URL or data-api-endpoint
            api_endpoint = img.get('data-api-endpoint', '')
            file_id = None

            # Try to extract from src URL
            file_id_match = re.search(r'/files/(\d+)', orig_src)
            if file_id_match:
                file_id = file_id_match.group(1)
            else:
                # Try data-api-endpoint
                file_id_match = re.search(r'/files/(\d+)', api_endpoint)
                if file_id_match:
                    file_id = file_id_match.group(1)

            # If we have a Canvas file ID, look for existing local file
            if file_id:
                # Determine the image directory based on HTML file location
                # Resolve to absolute paths first
                html_file_resolved = html_file.resolve()
                html_dir_resolved = html_dir.resolve()

                # In the new structure, images are in the same directory as the HTML file
                # Check for image in the same directory first
                possible_extensions = ['.png', '.jpg', '.jpeg', '.gif']
                found = False

                # First, check in the same directory as HTML file
                for ext in possible_extensions:
                    image_path = html_dir_resolved / f"{file_id}{ext}"
                    if image_path.exists():
                        img['src'] = f"{file_id}{ext}"
                        print(f"  ✅ Found local image {file_id}{ext} in same directory")
                        found = True
                        break

                # If not found, check in archive 365_Update directory (for backward compatibility)
                if not found:
                    # Try archive directory
                    archive_dir = get_winter_archive_dir() / "365_Update"
                    base_dir = None  # Initialize for use in except block
                    if archive_dir.exists():
                        image_dir = archive_dir
                        base_dir = get_winter_archive_dir()  # Set base_dir for path calculation
                    else:
                        # Fallback: search from HTML file location
                        base_dir = html_file_resolved.parent
                        while base_dir.parent != base_dir:
                            image_dir = base_dir / "365_Update"
                            if image_dir.exists():
                                break
                            base_dir = base_dir.parent
                        else:
                            image_dir = None

                    if image_dir and image_dir.exists():
                        for ext in possible_extensions:
                            # Search recursively in 365_Update
                            image_files = list(image_dir.rglob(f"{file_id}{ext}"))
                            if image_files:
                                image_path = image_files[0]
                                try:
                                    rel_path = image_path.resolve().relative_to(html_dir_resolved)
                                    img['src'] = str(rel_path)
                                    print(f"  ✅ Found local image {file_id}{ext} at: {rel_path}")
                                    found = True
                                    break
                                except ValueError:
                                    # Calculate relative path manually
                                    if base_dir:
                                        depth = len(html_dir_resolved.relative_to(base_dir).parts)
                                        rel_path = "../" * depth + str(image_path.relative_to(base_dir))
                                    else:
                                        rel_path = str(image_path.name)  # Fallback to just filename
                                    img['src'] = rel_path
                                    print(f"  ✅ Found local image {file_id}{ext} at: {rel_path}")
                                    found = True
                                    break

                if not found:
                    print(f"  ❌ Local image not found for Canvas ID {file_id}")
                    # Don't replace with "image not found" - keep the original URL
                    # img.replace_with("image not found")
                continue

            # Check if it's already a relative path (our new format)
            if orig_src.startswith('../') or orig_src.startswith('./'):
                # It's a relative path - verify the image exists
                image_path = (html_dir / orig_src).resolve()
                if image_path.exists():
                    # Keep the relative path - Pandoc will resolve it from HTML file location
                    print(f"  ✅ Found image at relative path: {orig_src}")
                    continue
                else:
                    print(f"  ❌ Image not found at: {image_path}")
                    img.replace_with("image not found")
            else:
                # Canvas URL but no file ID found, or other format
                print(f"  ❌ No Canvas file ID found for image: {orig_src[:60]}...")
                img.replace_with("image not found")

    # Clean up structure: unwrap empty divs and fix nested lists
    # First pass: unwrap empty divs
    for div in cleaned.find_all('div'):
        # If div has no meaningful content (only whitespace or other empty divs), unwrap it
        text_content = div.get_text(strip=True)
        child_divs = div.find_all('div', recursive=False)
        other_children = [c for c in div.children if c.name != 'div' and (hasattr(c, 'get_text') and c.get_text(strip=True))]

        if not text_content and not other_children and len(child_divs) == 1:
            # Unwrap single empty div wrapper
            div.unwrap()
        elif not text_content and not other_children and len(child_divs) > 1:
            # Multiple empty divs - unwrap the parent
            for child in list(div.children):
                if child.name == 'div':
                    div.parent.insert(div.parent.index(div), child)
            div.decompose()

    # Fix nested lists (ol within ol, ul within ul)
    for list_tag in cleaned.find_all(['ol', 'ul']):
        # Find nested lists of the same type
        nested_same = list_tag.find_all(list_tag.name, recursive=False)
        for nested in nested_same:
            # Unwrap nested list - move its items to parent
            for item in nested.find_all('li', recursive=False):
                list_tag.append(item)
            nested.decompose()

    # Map paragraphs and lists after headings to Word styles based on heading level
    # Style mapping:
    #   Heading 1 -> Body Text, List Bullet, List Number
    #   Heading 2 -> Body Text 2, List Bullet 2, List Number 2
    #   Heading 3 -> Body Text 3, List Bullet 3, List Number 3
    #   Heading 4+ -> Body Text 3, List Bullet 3, List Number 3 (use level 3 for deeper headings)

    # Get all headings in document order
    headings = cleaned.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])

    # Process each heading and apply styles to following content
    for i, heading in enumerate(headings):
        heading_level = int(heading.name[1]) if heading.name[1].isdigit() else 0

        # Determine style level (cap at 3)
        style_level = min(heading_level, 3)

        # Map heading level to style suffix
        if style_level == 1:
            para_style = 'Body Text'
            bullet_style = 'List Bullet'
            number_style = 'List Number'
        elif style_level == 2:
            para_style = 'Body Text 2'
            bullet_style = 'List Bullet 2'
            number_style = 'List Number 2'
        else:  # level 3 or higher
            para_style = 'Body Text 3'
            bullet_style = 'List Bullet 3'
            number_style = 'List Number 3'

        # Find the next heading (if any) to know where to stop
        next_heading = headings[i + 1] if i + 1 < len(headings) else None

        # Find all paragraphs and lists that come after this heading
        # Use find_all_next to get all following elements regardless of nesting
        following_elements = heading.find_all_next(['p', 'ul', 'ol', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'])

        # Process each following element until we hit the next heading
        for element in following_elements:
            # Stop if we hit the next heading
            if next_heading and element == next_heading:
                break

            # Also stop if we hit any heading at same or higher level
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                element_level = int(element.name[1]) if element.name[1].isdigit() else 0
                if element_level <= heading_level:
                    break

            # Apply styles to paragraphs and lists
            if element.name in ['p', 'ul', 'ol']:
                # Check if element is inside a callout/note/important/instructions/question div
                parent = element.parent
                is_in_special_div = False
                while parent and parent.name != 'body':
                    if parent.name == 'div' and 'class' in parent.attrs:
                        classes = parent.attrs['class']
                        if isinstance(classes, list):
                            if any(cls in ['Callout', 'Note', 'Important', 'Instructions', 'Question'] for cls in classes):
                                is_in_special_div = True
                                break
                    parent = parent.parent

                # Apply styles based on element type and heading level
                # Pandoc uses 'custom-style' attribute (not 'class') for Word paragraph styles
                if not is_in_special_div:
                    if element.name == 'p':
                        element.attrs['custom-style'] = para_style
                        # Remove class if it exists to avoid conflicts
                        if 'class' in element.attrs:
                            del element.attrs['class']
                    elif element.name == 'ul':
                        element.attrs['custom-style'] = bullet_style
                        if 'class' in element.attrs:
                            del element.attrs['class']
                    elif element.name == 'ol':
                        element.attrs['custom-style'] = number_style
                        if 'class' in element.attrs:
                            del element.attrs['class']

    # Map CSS classes to Word styles and clean up attributes
    # IMPORTANT: Preserve custom-style attributes that were set above
    for tag in cleaned.find_all(True):
        # Remove Canvas-specific data attributes (but preserve custom-style)
        for attr in list(tag.attrs.keys()):
            if attr.startswith('data-') and attr != 'custom-style':
                del tag.attrs[attr]

        # Handle guide div - keep it but remove class (styling handled by reference DOCX)
        if tag.name == 'div' and 'class' in tag.attrs:
            classes = tag.attrs['class']
            if isinstance(classes, list) and 'guide' in classes:
                # Keep guide div but remove class (Pandoc will treat as regular div)
                classes.remove('guide')
                if not classes:
                    del tag.attrs['class']
                else:
                    tag.attrs['class'] = classes

        # Preserve original-link div and ensure links are styled (blue and underlined)
        if tag.name == 'div' and 'class' in tag.attrs:
            classes = tag.attrs['class']
            if isinstance(classes, list) and 'original-link' in classes:
                # Keep the original-link div - don't remove it
                # Find all links in the original-link div and style them
                for link in tag.find_all('a'):
                    # Ensure href is preserved (Pandoc needs this for hyperlinks)
                    if not link.get('href'):
                        # Try to extract from existing attributes or keep as placeholder
                        pass
                    # Add inline styles for blue color and underline
                    # Use both style attribute and ensure Pandoc recognizes it as a hyperlink
                    existing_style = link.get('style', '')
                    style_parts = []
                    if existing_style:
                        style_parts.append(existing_style)
                    # Add blue color and underline
                    style_parts.append('color: #0000FF;')
                    style_parts.append('text-decoration: underline;')
                    link['style'] = ' '.join(style_parts)

        # Map CSS classes to Word styles for callouts
        # Use Pandoc's custom-style attribute for Word styles
        if 'class' in tag.attrs:
            classes = tag.attrs['class']
            if isinstance(classes, list):
                # Map callout classes to Word styles using Pandoc custom-style attribute
                # Style names must match Word style names in reference DOCX exactly
                if 'callout' in classes and 'instructions' in classes:
                    # Map to Instructions style
                    if tag.name == 'div':
                        tag.attrs['custom-style'] = 'Instructions'
                        tag.attrs['class'] = [c for c in classes if c not in ['callout', 'instructions']]
                        if not tag.attrs['class']:
                            del tag.attrs['class']
                elif 'callout' in classes:
                    # Map to Callout style
                    if tag.name == 'div':
                        tag.attrs['custom-style'] = 'Callout'
                        tag.attrs['class'] = [c for c in classes if c != 'callout']
                        if not tag.attrs['class']:
                            del tag.attrs['class']
                elif 'note' in classes:
                    # Map to Note style
                    if tag.name == 'div':
                        tag.attrs['custom-style'] = 'Note'
                        tag.attrs['class'] = [c for c in classes if c != 'note']
                        if not tag.attrs['class']:
                            del tag.attrs['class']
                elif 'important' in classes:
                    # Map to Important style
                    if tag.name == 'div':
                        tag.attrs['custom-style'] = 'Important'
                        tag.attrs['class'] = [c for c in classes if c != 'important']
                        if not tag.attrs['class']:
                            del tag.attrs['class']
                elif 'instructions' in classes:
                    # Map to Instructions style
                    if tag.name == 'div':
                        tag.attrs['custom-style'] = 'Instructions'
                        tag.attrs['class'] = [c for c in classes if c != 'instructions']
                        if not tag.attrs['class']:
                            del tag.attrs['class']
                elif 'question2' in classes or 'questionBullet' in classes:
                    # Map to Question style
                    if tag.name == 'div':
                        tag.attrs['custom-style'] = 'Question'
                        tag.attrs['class'] = [c for c in classes if c not in ['question2', 'questionBullet']]
                        if not tag.attrs['class']:
                            del tag.attrs['class']

                # Remove Canvas UI classes but keep Word style classes
                if 'class' in tag.attrs:
                    # Keep Word style classes that match reference DOCX styles
                    word_style_classes = ['Callout', 'Note', 'Important', 'Instructions', 'Question',
                                        'Body Text', 'Body Text 2', 'Body Text 3',
                                        'List', 'List 2', 'List 3',
                                        'List Bullet', 'List Bullet 2', 'List Bullet 3',
                                        'List Number', 'List Number 2', 'List Number 3',
                                        'List Continue', 'List Continue 2', 'List Continue 3',
                                        'List Paragraph', 'Caption', 'Quote', 'Intense Quote']
                    filtered_classes = [c for c in tag.attrs['class']
                                      if c not in ['lti-embed', 'inline_disabled', 'screen', 's2', 'guide',
                                                   'callout', 'note', 'important', 'instructions', 'question2', 'questionBullet',
                                                   'h3-paragraph', 'h4-paragraph', 'h5-paragraph',
                                                   'paragraph-text-after-a-heading', 'h5-numbers-list']
                                      or c in word_style_classes]
                    if filtered_classes:
                        tag.attrs['class'] = filtered_classes
                    else:
                        del tag.attrs['class']

        # Handle images with screen class - add border styling hint
        if tag.name == 'img' and 'class' in tag.attrs:
            classes = tag.attrs['class']
            if isinstance(classes, list) and 'screen' in classes:
                # Add border style hint (Pandoc may not fully support, but worth trying)
                if 'style' not in tag.attrs:
                    tag['style'] = 'border: 1px solid #CCCCCC; border-radius: 4px; padding: 0.5em;'
                classes.remove('screen')
                if not classes:
                    del tag.attrs['class']
                else:
                    tag.attrs['class'] = classes

        # Preserve some style attributes that affect formatting
        # (Pandoc can handle some CSS)
        # IMPORTANT: Always preserve styles on links (especially in original-link div)
        if 'style' in tag.attrs:
            style = tag.attrs['style']
            # Always keep styles on anchor tags (for hyperlink formatting)
            if tag.name == 'a':
                # Keep all styles on links - Pandoc will use them
                pass
            # Keep important formatting styles on other elements
            elif 'font-weight' in style or 'font-style' in style or 'text-align' in style or 'border' in style or 'color' in style or 'text-decoration' in style:
                # Keep it - Pandoc might use it
                pass
            else:
                # Remove layout styles that won't translate
                del tag.attrs['style']

    # Convert <br/> tags to newlines in paragraphs (but keep them for Pandoc)
    # Actually, let Pandoc handle <br/> tags - they work in HTML

    return cleaned

def convert_html_to_docx(html_content, output_docx_path, html_file_path=None, reference_doc=None):
    """Convert HTML content to DOCX using Pandoc."""
    # Check if Pandoc is installed
    try:
        result = subprocess.run(['pandoc', '--version'],
                              capture_output=True, text=True)
        if result.returncode != 0:
            raise FileNotFoundError("Pandoc not found")
    except FileNotFoundError:
        print("❌ Error: Pandoc is not installed.")
        print("   Install with: brew install pandoc")
        print("   Or download from: https://pandoc.org/installing.html")
        sys.exit(1)

    # Determine working directory for Pandoc
    # Run from HTML file's directory so relative image paths resolve correctly
    if html_file_path:
        work_dir = Path(html_file_path).parent
    else:
        work_dir = output_docx_path.parent

    # Write HTML to temporary file in the working directory
    temp_html = work_dir / f"{output_docx_path.stem}.temp.html"
    with open(temp_html, 'w', encoding='utf-8') as f:
        f.write(str(html_content))

    # Build Pandoc command with options to preserve formatting and apply styles
    # Pandoc automatically maps HTML elements and classes to Word styles from reference DOCX:
    # - h1-h6 -> Heading 1-6
    # - div with class="StyleName" -> applies Word style "StyleName" if it exists in reference DOCX
    # - p, ul, ol -> Normal, List Bullet, List Number (or custom styles if classes match)
    cmd = [
        'pandoc',
        temp_html.name,  # relative to work_dir
        '-o', str(output_docx_path.resolve()),  # absolute path for output
        '--from', 'html',
        '--to', 'docx',
        '--standalone',  # Include header/footer
        '--wrap=none',   # Don't wrap lines
        '--preserve-tabs',  # Preserve tabs
    ]

    # Add reference document if provided
    if reference_doc and Path(reference_doc).exists():
        ref_doc_path = Path(reference_doc).resolve()
        cmd.extend(['--reference-doc', str(ref_doc_path)])
        print(f"   Using reference DOCX: {ref_doc_path}")
    else:
        print(f"   ⚠️  Reference DOCX not found, using default styles")

    # Run Pandoc from the HTML file's directory so relative image paths resolve
    print(f"🔄 Converting HTML to DOCX using Pandoc...")
    print(f"   Working directory: {work_dir}")
    print(f"   Output: {output_docx_path}")
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=work_dir)

    # Clean up temp file
    temp_html.unlink(missing_ok=True)

    if result.returncode != 0:
        print(f"❌ Pandoc conversion failed:")
        print(result.stderr)
        sys.exit(1)

    print(f"✅ DOCX file created: {output_docx_path}")

def parse_canvas_urls_from_course_map(course_map_path):
    """Parse course-map.md and extract Canvas URLs mapped by page slug.

    Returns a dict mapping page slugs (normalized) to Canvas URLs.
    """
    import re

    url_map = {}

    if not Path(course_map_path).exists():
        return url_map

    with open(course_map_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Pattern to match Canvas URL lines
    # Format: - **Canvas URL:** https://usucourses.instructure.com/courses/2879/pages/{slug}
    pattern = r'\*\*Canvas URL:\*\*\s+(https://usucourses\.instructure\.com/courses/2879/pages/[^\s]+)'

    for match in re.finditer(pattern, content):
        canvas_url = match.group(1)
        # Extract slug from URL
        slug_match = re.search(r'/pages/([^\s]+)$', canvas_url)
        if slug_match:
            slug = slug_match.group(1)
            # Store with slug as key (normalized)
            url_map[slug] = canvas_url

    return url_map

def get_canvas_url_from_html_file(html_file_path, course_map_path=None):
    """Get Canvas URL from HTML file path.

    First tries to get URL from course-map.md if available.
    Falls back to deriving URL from filename if course map lookup fails.
    """
    html_file = Path(html_file_path)

    # Try course map first if path provided
    if course_map_path:
        course_map_file = Path(course_map_path)
        if not course_map_file.exists():
            # Try relative to project root
            project_root = get_project_root()
            course_map_file = project_root / 'data' / 'current' / 'WINTER-25-26-UPDATES' / 'course-map.md'

        if course_map_file.exists():
            url_map = parse_canvas_urls_from_course_map(course_map_file)

            # Get base name without extension for lookup
            base_name = html_file.stem.lower()

            # Remove '_local' suffix if present
            if base_name.endswith('_local'):
                base_name = base_name[:-6]

            # Try exact match first (slug format: underscores to hyphens)
            slug = base_name.replace('_', '-')
            if slug in url_map:
                return url_map[slug]

            # Try with different variations (some slugs might have different formatting)
            # Check if there's a match with variations
            for map_slug, url in url_map.items():
                # Check if the slug ends with our base slug or vice versa
                if map_slug.endswith(slug) or slug.endswith(map_slug) or slug in map_slug or map_slug in slug:
                    return url

    # Fallback: derive Canvas URL from filename
    base_name = html_file.stem.lower()

    # Remove '_local' suffix if present
    if base_name.endswith('_local'):
        base_name = base_name[:-6]

    # Replace underscores with hyphens to get Canvas page slug
    slug = base_name.replace('_', '-')

    # Construct Canvas URL
    canvas_url = f"https://usucourses.instructure.com/courses/2879/pages/{slug}"

    return canvas_url

def add_canvas_link_to_docx(docx_path, canvas_url):
    """Add 'View original page on Canvas' hyperlink at the top of DOCX file."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document(str(docx_path))

    # Check if link already exists
    link_exists = False
    for para in doc.paragraphs[:3]:  # Check first 3 paragraphs
        if 'View original page' in para.text:
            link_exists = True
            break

    if link_exists:
        print(f"  ℹ️  Canvas link already exists in document")
        return

    # Create new paragraph for the link at the beginning
    link_para = doc.paragraphs[0].insert_paragraph_before()
    link_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add the hyperlink
    # First, add the relationship
    part = doc.part
    r_id = part.relate_to(canvas_url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create run element with text
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # Set underline
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    # Add text
    text = OxmlElement('w:t')
    text.text = 'View original page on Canvas'
    text.set(qn('xml:space'), 'preserve')

    run.append(rPr)
    run.append(text)
    hyperlink.append(run)

    # Add hyperlink to paragraph
    link_para._p.append(hyperlink)

    # Apply FirstParagraph style if it exists, otherwise use Normal
    try:
        link_para.style = 'FirstParagraph'
    except:
        link_para.style = 'Normal'

    doc.save(str(docx_path))
    print(f"  ✅ Added Canvas link at top: {canvas_url}")

def main():
    parser = argparse.ArgumentParser(
        description='Convert Canvas HTML page to DOCX file'
    )
    parser.add_argument(
        '--html-file',
        type=Path,
        required=True,
        help='Path to Canvas HTML file'
    )
    parser.add_argument(
        '--output-docx',
        type=Path,
        help='Output DOCX file path (default: same directory as HTML with .docx extension)'
    )
    parser.add_argument(
        '--reference-doc',
        type=Path,
        default=get_canvas_reference_docx(),
        help='Reference DOCX template for styling (default: assets/templates/canvas-reference.docx)'
    )

    args = parser.parse_args()

    # Set default output path if not provided
    if not args.output_docx:
        # Generate filename: lowercase, spaces to underscores, remove "_local" suffix
        base_name = args.html_file.stem.lower().replace(' ', '_')
        # Remove "_local" suffix if present (e.g., "course_orientation_local" -> "course_orientation")
        if base_name.endswith('_local'):
            base_name = base_name[:-6]  # Remove "_local" (6 characters)
        args.output_docx = args.html_file.parent / f"{base_name}.docx"

    # Extract user_content from HTML (no downloads, only use existing local files with Canvas IDs)
    print(f"📄 Extracting content from: {args.html_file}")
    html_content = extract_user_content(args.html_file)
    print(f"✅ Extracted user_content div")

    # Convert to DOCX
    convert_html_to_docx(html_content, args.output_docx, args.html_file, args.reference_doc)

    # Add "View original page on Canvas" link at the top
    print(f"\n🔗 Adding Canvas link...")
    try:
        # Try to use course-map.md if available
        project_root = get_project_root()
        course_map_path = project_root / 'data' / 'current' / 'WINTER-25-26-UPDATES' / 'course-map.md'
        canvas_url = get_canvas_url_from_html_file(args.html_file, course_map_path)
        add_canvas_link_to_docx(args.output_docx, canvas_url)
    except Exception as e:
        print(f"  ⚠️  Could not add Canvas link: {e}")
        import traceback
        traceback.print_exc()

    # Enable Track Changes mode so document opens in "Reviewing" mode in Word Online
    print(f"\n🔧 Enabling Track Changes mode...")
    try:
        import subprocess
        # enable-track-changes.py is in scripts/utility/
        enable_script = Path(__file__).parent.parent / 'utility' / 'enable-track-changes.py'
        result = subprocess.run(
            ['python3', str(enable_script), str(args.output_docx)],
            capture_output=True,
            text=True
        )
        if result.returncode == 0:
            print(result.stdout)
        else:
            print(f"  ⚠️  Could not enable Track Changes: {result.stderr}")
    except Exception as e:
        print(f"  ⚠️  Could not enable Track Changes: {e}")

    # Style the "View original page on Canvas" hyperlink (blue and underlined)
    # This ensures any existing links are properly styled, though our add_canvas_link_to_docx
    # function already creates them with the correct styling
    print(f"\n🔗 Verifying Canvas hyperlink styling...")
    try:
        from docx import Document
        from docx.shared import RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document(str(args.output_docx))

        # Find the paragraph with "View original page on Canvas"
        for para in doc.paragraphs:
            if 'View original page' in para.text:
                # Find all hyperlink elements in this paragraph
                # Use findall with namespace-aware search
                hyperlinks = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')

                if hyperlinks:
                    for hyperlink in hyperlinks:
                        # Find all runs inside this hyperlink
                        runs = hyperlink.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')

                        for run_elem in runs:
                            # Get or create rPr (run properties) element
                            rPr = run_elem.find(qn('w:rPr'))
                            if rPr is None:
                                rPr = OxmlElement('w:rPr')
                                run_elem.insert(0, rPr)

                            # Remove rStyle if it exists (it might override direct formatting)
                            rStyle = rPr.find(qn('w:rStyle'))
                            if rStyle is not None:
                                rPr.remove(rStyle)

                            # Add blue color (ensure it's set, not just appended)
                            color_elem = rPr.find(qn('w:color'))
                            if color_elem is None:
                                color_elem = OxmlElement('w:color')
                                color_elem.set(qn('w:val'), '0000FF')  # Blue in hex
                                rPr.append(color_elem)
                            else:
                                color_elem.set(qn('w:val'), '0000FF')

                            # Add underline (ensure it's set, not just appended)
                            u_elem = rPr.find(qn('w:u'))
                            if u_elem is None:
                                u_elem = OxmlElement('w:u')
                                u_elem.set(qn('w:val'), 'single')
                                rPr.append(u_elem)
                            else:
                                u_elem.set(qn('w:val'), 'single')

                            # Get text from the run
                            text_elems = run_elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                            text = ''.join([t.text for t in text_elems if t.text])
                            print(f"  ✅ Verified hyperlink styling: {text[:40]}")
                else:
                    # If no hyperlink found, try to style runs directly
                    for run in para.runs:
                        if 'View original page' in run.text:
                            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                            run.font.underline = True
                            print(f"  ✅ Styled text: {run.text[:40]}")

                break

        doc.save(str(args.output_docx))
        print(f"  ✅ Hyperlink styling verified")
    except Exception as e:
        print(f"  ⚠️  Could not verify hyperlink styling: {e}")
        import traceback
        traceback.print_exc()

    print(f"\n📋 Next steps:")
    print(f"   1. Open {args.output_docx} in Word")
    print(f"   2. Review and adjust formatting if needed")
    print(f"   3. Upload to Box (replace existing DOCX)")
    print(f"   4. Document will open in 'Reviewing' mode with Track Changes enabled")
    print(f"   5. Run create-docx-html-mapping.py to create new mapping")
    print(f"   6. Test with a tracked change")

if __name__ == '__main__':
    main()

